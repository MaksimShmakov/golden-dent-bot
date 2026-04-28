"""Конвертер Markdown -> .docx для документов проекта.

По умолчанию собирает ARCHITECTURE.md -> ARCHITECTURE.docx. Можно
передать произвольный путь к .md-файлу аргументом командной строки;
итоговый .docx будет сохранён рядом с тем же базовым именем.

Поддерживает ограниченное подмножество Markdown: заголовки #/##/###/####,
маркированные списки "- ", нумерованные списки "N.", таблицы в формате
GFM, жирный **текст**, инлайн `code`, горизонтальные разделители "---".
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
DEFAULT_SOURCE = ROOT / "ARCHITECTURE.md"

BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
INLINE_CODE_RE = re.compile(r"`([^`]+)`")


def add_runs(paragraph, text: str) -> None:
    """Добавляет текст с инлайн-форматированием (жирный, код)."""
    pos = 0
    pattern = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        chunk = match.group(0)
        if chunk.startswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(chunk[1:-1])
            run.font.name = "Consolas"
            run.font.color.rgb = RGBColor(0x6A, 0x1B, 0x9A)
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def set_base_style(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for section in document.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)


def render_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = len(rows[0])
    table = document.add_table(rows=len(rows), cols=cols)
    table.style = "Light Grid Accent 1"
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx].cells
        for c_idx in range(cols):
            cell = cells[c_idx]
            cell.text = ""
            para = cell.paragraphs[0]
            add_runs(para, row[c_idx] if c_idx < len(row) else "")
            if r_idx == 0:
                for run in para.runs:
                    run.bold = True


def convert(md_text: str, document: Document) -> None:
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip():
            i += 1
            continue

        if line.strip() == "---":
            document.add_paragraph().add_run().add_break()
            i += 1
            continue

        if line.startswith("# "):
            heading = document.add_heading(line[2:].strip(), level=0)
            heading.paragraph_format.space_after = Pt(12)
            i += 1
            continue
        if line.startswith("## "):
            document.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        if line.startswith("### "):
            document.add_heading(line[4:].strip(), level=2)
            i += 1
            continue
        if line.startswith("#### "):
            document.add_heading(line[5:].strip(), level=3)
            i += 1
            continue

        # Таблица GFM: строка с | ... | и следующая строка — разделитель
        if line.lstrip().startswith("|") and i + 1 < len(lines) and re.match(r"\s*\|[\s:\-|]+\|\s*$", lines[i + 1]):
            rows: list[list[str]] = []
            header = [c.strip() for c in line.strip().strip("|").split("|")]
            rows.append(header)
            i += 2  # пропускаем разделитель
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(row)
                i += 1
            render_table(document, rows)
            continue

        # Маркированный список
        if re.match(r"\s*-\s+", line):
            while i < len(lines) and re.match(r"\s*-\s+", lines[i]):
                stripped = lines[i].lstrip()
                indent = len(lines[i]) - len(stripped)
                text = stripped[2:]
                style = "List Bullet" if indent < 2 else "List Bullet 2"
                para = document.add_paragraph(style=style)
                add_runs(para, text)
                i += 1
            continue

        # Нумерованный список
        if re.match(r"\s*\d+\.\s+", line):
            while i < len(lines) and re.match(r"\s*\d+\.\s+", lines[i]):
                stripped = lines[i].lstrip()
                text = re.sub(r"^\d+\.\s+", "", stripped)
                para = document.add_paragraph(style="List Number")
                add_runs(para, text)
                i += 1
            continue

        # Обычный абзац
        para = document.add_paragraph()
        add_runs(para, line)
        i += 1


def main(argv: list[str] | None = None) -> None:
    args = argv if argv is not None else sys.argv[1:]
    source = Path(args[0]).resolve() if args else DEFAULT_SOURCE
    target = source.with_suffix(".docx")
    md_text = source.read_text(encoding="utf-8")
    document = Document()
    set_base_style(document)
    convert(md_text, document)
    document.save(target)
    print(f"Saved: {target}")


if __name__ == "__main__":
    main()
